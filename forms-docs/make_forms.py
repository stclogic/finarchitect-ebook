# 표준 서식 DOCX 일괄 생성 스크립트
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "/Users/stclogic/Desktop/stclogic-ebook/forms-docs/"

# ── 공통 유틸 ──────────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def title_para(doc, text, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    return p

def heading(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=(30, 80, 160))
    return p

def body(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def blank_line(doc, text="    "):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.space_after = Pt(2)
    return p

def field_line(doc, label, lines=1):
    """라벨: _____ 형식 항목"""
    p = doc.add_paragraph()
    run = p.add_run(f"  {label}:  " + "_" * 40)
    set_font(run, size=11)
    return p

def divider(doc):
    p = doc.add_paragraph("─" * 60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    set_font(run, size=9, color=(150, 150, 150))

def sig_table(doc, parties):
    """서명란 테이블"""
    doc.add_paragraph()
    table = doc.add_table(rows=len(parties), cols=3)
    table.style = "Table Grid"
    headers = ["구분", "성명 (서명)", "날짜"]
    for i, party in enumerate(parties):
        row = table.rows[i]
        row.cells[0].text = party
        row.cells[1].text = "                    (인)"
        row.cells[2].text = "____년 __월 __일"
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, size=10)
    return table

def page_setup(doc):
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

def watermark_note(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("본 서식은 STCLOGIC (finarchitect.com) 에서 제공하는 표준 양식입니다.")
    set_font(run, size=8, color=(180, 180, 180))


# ══════════════════════════════════════════════════════════════════
# 1. 근로계약서 (정규직)
# ══════════════════════════════════════════════════════════════════
def make_labor_contract():
    doc = Document()
    page_setup(doc)

    title_para(doc, "근 로 계 약 서", 18)
    title_para(doc, "( 정규직 )", 12)
    doc.add_paragraph()

    body(doc, "사용자와 근로자는 근로기준법에 따라 아래와 같이 근로계약을 체결한다.")
    doc.add_paragraph()

    heading(doc, "■ 당사자 정보")
    field_line(doc, "사용자 (회사명)")
    field_line(doc, "사업자등록번호")
    field_line(doc, "주 소")
    field_line(doc, "대표자")
    doc.add_paragraph()
    field_line(doc, "근로자 성명")
    field_line(doc, "주민등록번호")
    field_line(doc, "주 소")
    field_line(doc, "연락처")
    doc.add_paragraph()

    heading(doc, "■ 제1조 (근무 장소)")
    body(doc, "근무 장소:  " + "_" * 50)
    doc.add_paragraph()

    heading(doc, "■ 제2조 (업무 내용)")
    body(doc, "담당 업무:  " + "_" * 50)
    doc.add_paragraph()

    heading(doc, "■ 제3조 (근로 기간)")
    body(doc, "입사일:  ____년 __월 __일  (정규직 — 기간의 정함 없음)")
    doc.add_paragraph()

    heading(doc, "■ 제4조 (근로 시간)")
    body(doc, "소정 근로 시간: 09:00 ~ 18:00  (휴게 시간: 12:00 ~ 13:00, 1시간)")
    body(doc, "주 40시간 / 주 5일 근무 (월요일 ~ 금요일)")
    doc.add_paragraph()

    heading(doc, "■ 제5조 (임금)")
    field_line(doc, "월 기본급 (세전)")
    field_line(doc, "상여금 및 기타 수당")
    body(doc, "임금 지급일: 매월 ___일  /  지급 방법: 근로자 명의 계좌 이체")
    doc.add_paragraph()

    heading(doc, "■ 제6조 (연차 유급 휴가)")
    body(doc, "근로기준법 제60조에 따라 연차 유급 휴가를 부여한다.")
    doc.add_paragraph()

    heading(doc, "■ 제7조 (사회보험)")
    body(doc, "4대보험(국민연금·건강보험·고용보험·산재보험)을 관계 법령에 따라 적용한다.")
    doc.add_paragraph()

    heading(doc, "■ 제8조 (기타)")
    body(doc, "본 계약에 명시되지 않은 사항은 근로기준법 및 취업 규칙에 따른다.")
    doc.add_paragraph()

    divider(doc)
    body(doc, "본 계약서는 2부 작성하여 사용자와 근로자가 각 1부씩 보관한다.")
    doc.add_paragraph()
    body(doc, "계약 체결일:  ____년 __월 __일")
    doc.add_paragraph()

    sig_table(doc, ["사용자 (회사)", "근로자"])
    doc.add_paragraph()
    watermark_note(doc)

    doc.save(OUTPUT + "01_근로계약서_정규직.docx")
    print("✓ 01_근로계약서_정규직.docx")


# ══════════════════════════════════════════════════════════════════
# 2. 표준 임대차계약서
# ══════════════════════════════════════════════════════════════════
def make_lease_contract():
    doc = Document()
    page_setup(doc)

    title_para(doc, "부동산 임대차계약서", 18)
    title_para(doc, "( 주택 / 상가 )", 11)
    doc.add_paragraph()

    heading(doc, "■ 부동산의 표시")
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("소재지", ""),
        ("토지", "지목:         면적:           ㎡"),
        ("건물", "구조/용도:         면적:           ㎡"),
        ("임대 부분", "층:     호수:     면적:           ㎡"),
        ("계약 유형", "□ 전세   □ 월세   □ 반전세"),
    ]
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=10)
    doc.add_paragraph()

    heading(doc, "■ 계약 내용")
    body(doc, "보증금:  금  _____________원  정  (₩               )")
    body(doc, "월 차임:  금  _____________원  정  (₩               )  — 매월 ___일 지급")
    body(doc, "임대 기간:  ____년 __월 __일  ~  ____년 __월 __일  (___개월)")
    doc.add_paragraph()

    heading(doc, "■ 제1조 (목적)")
    body(doc, "임대인은 위 부동산을 임차인에게 임대하고, 임차인은 이를 임차하여 사용·수익하기로 한다.")
    doc.add_paragraph()

    heading(doc, "■ 제2조 (보증금)")
    body(doc, "임차인은 계약 체결 시 계약금, 잔금은 입주일에 각각 지급한다.")
    body(doc, "계약금:  금  _____________원  (계약 당일 지급)")
    body(doc, "잔  금:  금  _____________원  (____년 __월 __일 지급)")
    doc.add_paragraph()

    heading(doc, "■ 제3조 (사용 목적)")
    body(doc, "임차인은 임대 부동산을 □ 주거용  □ 영업용  □ 기타( _____________ ) 목적으로만 사용한다.")
    doc.add_paragraph()

    heading(doc, "■ 제4조 (수선 및 유지)")
    body(doc, "임대인은 목적물을 임차인이 사용·수익할 수 있는 상태로 유지할 의무를 진다.")
    body(doc, "소모품 및 소액 수선(30만원 미만)은 임차인 부담으로 한다.")
    doc.add_paragraph()

    heading(doc, "■ 제5조 (계약 해지)")
    body(doc, "임대차 기간 만료 2개월 전까지 갱신 거절 통보가 없으면 동일 조건으로 묵시적 갱신된다.")
    body(doc, "임차인의 귀책으로 계약 해지 시, 임대인은 손해를 공제 후 보증금을 반환한다.")
    doc.add_paragraph()

    heading(doc, "■ 제6조 (특약 사항)")
    for _ in range(4):
        body(doc, "  ① " + "_" * 60)
    doc.add_paragraph()

    divider(doc)
    body(doc, "계약 체결일:  ____년 __월 __일")
    doc.add_paragraph()
    sig_table(doc, ["임대인", "임차인"])
    doc.add_paragraph()
    body(doc, "중개업소:  상호  _____________  대표  _____________  등록번호  _____________")
    doc.add_paragraph()
    watermark_note(doc)

    doc.save(OUTPUT + "02_임대차계약서.docx")
    print("✓ 02_임대차계약서.docx")


# ══════════════════════════════════════════════════════════════════
# 3. 차용증
# ══════════════════════════════════════════════════════════════════
def make_loan():
    doc = Document()
    page_setup(doc)

    title_para(doc, "차  용  증", 20)
    doc.add_paragraph()

    body(doc, "아래와 같이 금전을 차용하였음을 확인하고, 이를 증명하기 위하여 본 차용증을 작성합니다.")
    doc.add_paragraph()

    heading(doc, "■ 차용 내역")
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("차용금액", "금  _________________  원  정  (₩                    )"),
        ("이자율", "연   _____  %  (월   _____  %)  / □ 무이자"),
        ("이자 지급일", "매월   ____  일"),
        ("변제 기일", "____년   __월   __일"),
        ("변제 방법", "□ 일시 변제   □ 분할 변제  ( 매월   _____   원씩 )"),
        ("변제 계좌", "은행:              계좌번호:                  예금주:"),
    ]
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=10)
    doc.add_paragraph()

    heading(doc, "■ 당사자")
    field_line(doc, "채권자 (빌려주는 사람) 성명")
    field_line(doc, "주민등록번호")
    field_line(doc, "주소")
    field_line(doc, "연락처")
    doc.add_paragraph()
    field_line(doc, "채무자 (빌리는 사람) 성명")
    field_line(doc, "주민등록번호")
    field_line(doc, "주소")
    field_line(doc, "연락처")
    doc.add_paragraph()

    heading(doc, "■ 특약 사항")
    body(doc, "① 채무자가 변제 기일까지 원금 및 이자를 이행하지 않을 경우, 연 ___% 의 지연손해금을 부담한다.")
    body(doc, "② 분쟁 발생 시 채권자 주소지 관할 법원을 합의 관할로 한다.")
    for _ in range(2):
        body(doc, "③ " + "_" * 60)
    doc.add_paragraph()

    divider(doc)
    body(doc, "작성일:  ____년 __월 __일")
    doc.add_paragraph()
    sig_table(doc, ["채권자", "채무자"])
    doc.add_paragraph()
    watermark_note(doc)

    doc.save(OUTPUT + "03_차용증.docx")
    print("✓ 03_차용증.docx")


# ══════════════════════════════════════════════════════════════════
# 4. 비밀유지계약서 (NDA)
# ══════════════════════════════════════════════════════════════════
def make_nda():
    doc = Document()
    page_setup(doc)

    title_para(doc, "비밀유지계약서", 18)
    title_para(doc, "Non-Disclosure Agreement (NDA)", 11)
    doc.add_paragraph()

    heading(doc, "■ 당사자")
    field_line(doc, '제공자 (갑) 상호/성명')
    field_line(doc, "사업자등록번호 / 주민등록번호")
    field_line(doc, "주소")
    field_line(doc, "대표자 / 담당자")
    doc.add_paragraph()
    field_line(doc, '수령자 (을) 상호/성명')
    field_line(doc, "사업자등록번호 / 주민등록번호")
    field_line(doc, "주소")
    field_line(doc, "대표자 / 담당자")
    doc.add_paragraph()

    heading(doc, "■ 제1조 (목적)")
    body(doc, '갑과 을은 아래 사업과 관련하여 교류되는 비밀정보를 보호하기 위하여 본 계약을 체결한다.')
    field_line(doc, "협의 사업 내용")
    doc.add_paragraph()

    heading(doc, "■ 제2조 (비밀정보의 정의)")
    body(doc, '본 계약에서 "비밀정보"란 갑이 을에게 제공하는 기술, 경영, 재무, 고객 정보 등 일체를 말한다.')
    body(doc, "단, 다음의 경우는 비밀정보에서 제외한다:")
    body(doc, "  ① 공개 당시 이미 공지된 정보")
    body(doc, "  ② 수령 후 을의 귀책 없이 공지된 정보")
    body(doc, "  ③ 수령 전 을이 이미 보유한 정보")
    doc.add_paragraph()

    heading(doc, "■ 제3조 (비밀유지 의무)")
    body(doc, '을은 비밀정보를 목적 이외에 사용하거나 제3자에게 공개·누설하여서는 아니 된다.')
    body(doc, '을은 비밀정보 접근을 업무상 필요한 임직원으로 한정하고, 동등한 비밀유지 의무를 부과하여야 한다.')
    doc.add_paragraph()

    heading(doc, "■ 제4조 (비밀정보의 반환)")
    body(doc, '갑의 요청 또는 계약 종료 시, 을은 비밀정보 및 복사본을 즉시 반환하거나 파기하여야 한다.')
    doc.add_paragraph()

    heading(doc, "■ 제5조 (유효 기간)")
    body(doc, "본 계약의 유효 기간:  ____년 __월 __일  ~  ____년 __월 __일  (___개월)")
    body(doc, "비밀유지 의무는 계약 종료 후  ___  년간 지속된다.")
    doc.add_paragraph()

    heading(doc, "■ 제6조 (손해배상)")
    body(doc, '을이 본 계약을 위반하여 갑에게 손해를 끼친 경우 이를 배상할 책임을 진다.')
    doc.add_paragraph()

    heading(doc, "■ 제7조 (준거법 및 관할)")
    body(doc, "본 계약은 대한민국 법률에 따르며, 분쟁 시 갑 소재지 관할 법원을 제1심 관할 법원으로 한다.")
    doc.add_paragraph()

    divider(doc)
    body(doc, "계약 체결일:  ____년 __월 __일")
    doc.add_paragraph()
    sig_table(doc, ["제공자 (갑)", "수령자 (을)"])
    doc.add_paragraph()
    watermark_note(doc)

    doc.save(OUTPUT + "04_비밀유지계약서_NDA.docx")
    print("✓ 04_비밀유지계약서_NDA.docx")


# ══════════════════════════════════════════════════════════════════
# 5. 사업계획서 양식
# ══════════════════════════════════════════════════════════════════
def make_business_plan():
    doc = Document()
    page_setup(doc)

    title_para(doc, "사  업  계  획  서", 20)
    doc.add_paragraph()
    field_line(doc, "회사명 / 브랜드명")
    field_line(doc, "작성일")
    field_line(doc, "작성자")
    doc.add_paragraph()

    heading(doc, "■ 1. 사업 개요")
    body(doc, "사업명:  " + "_" * 50)
    body(doc, "사업 분야:  □ IT/플랫폼   □ 제조   □ 유통   □ 서비스   □ 기타 ( _______ )")
    body(doc, "핵심 한 줄 소개:")
    body(doc, "_" * 70)
    doc.add_paragraph()

    heading(doc, "■ 2. 문제 정의 (Problem)")
    for i in range(3):
        body(doc, f"  {'①②③'[i]} " + "_" * 60)
    doc.add_paragraph()

    heading(doc, "■ 3. 솔루션 (Solution)")
    for i in range(3):
        body(doc, f"  {'①②③'[i]} " + "_" * 60)
    doc.add_paragraph()

    heading(doc, "■ 4. 비즈니스 모델 (Revenue Model)")
    body(doc, "수익원:  □ 구독   □ 거래수수료   □ 광고   □ 제품판매   □ 기타 ( _______ )")
    body(doc, "가격 정책:")
    body(doc, "_" * 70)
    doc.add_paragraph()

    heading(doc, "■ 5. 목표 시장 (Target Market)")
    body(doc, "타겟 고객:  " + "_" * 50)
    body(doc, "시장 규모 (TAM/SAM/SOM):")
    body(doc, "_" * 70)
    doc.add_paragraph()

    heading(doc, "■ 6. 경쟁 분석")
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    headers = ["구분", "자사", "경쟁사 A", "경쟁사 B"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for r in table.rows[0].cells[i].paragraphs[0].runs:
            set_font(r, size=10, bold=True)
    for row_label in ["핵심 강점", "가격", "차별점"]:
        idx = ["핵심 강점","가격","차별점"].index(row_label) + 1
        table.rows[idx].cells[0].text = row_label
    doc.add_paragraph()

    heading(doc, "■ 7. 팀 구성")
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = "Table Grid"
    for i, h in enumerate(["성명", "직책", "주요 경력"]):
        table2.rows[0].cells[i].text = h
        for r in table2.rows[0].cells[i].paragraphs[0].runs:
            set_font(r, size=10, bold=True)
    doc.add_paragraph()

    heading(doc, "■ 8. 재무 계획 (향후 3년)")
    table3 = doc.add_table(rows=4, cols=4)
    table3.style = "Table Grid"
    for i, h in enumerate(["항목", "1년차", "2년차", "3년차"]):
        table3.rows[0].cells[i].text = h
        for r in table3.rows[0].cells[i].paragraphs[0].runs:
            set_font(r, size=10, bold=True)
    for j, label in enumerate(["매출", "영업이익", "투자 유치 목표"]):
        table3.rows[j+1].cells[0].text = label
    doc.add_paragraph()

    heading(doc, "■ 9. 투자 유치 계획")
    body(doc, "목표 투자금액:  " + "_" * 20 + "  원")
    body(doc, "자금 사용 계획:  □ 개발   □ 마케팅   □ 인건비   □ 운영   □ 기타")
    body(doc, "희망 투자 조건 (밸류에이션, 지분율 등):")
    body(doc, "_" * 70)
    doc.add_paragraph()

    heading(doc, "■ 10. 실행 로드맵")
    table4 = doc.add_table(rows=5, cols=3)
    table4.style = "Table Grid"
    for i, h in enumerate(["단계", "기간", "주요 목표"]):
        table4.rows[0].cells[i].text = h
        for r in table4.rows[0].cells[i].paragraphs[0].runs:
            set_font(r, size=10, bold=True)
    for j, label in enumerate(["Phase 1", "Phase 2", "Phase 3", "Phase 4"]):
        table4.rows[j+1].cells[0].text = label
    doc.add_paragraph()

    divider(doc)
    watermark_note(doc)

    doc.save(OUTPUT + "05_사업계획서.docx")
    print("✓ 05_사업계획서.docx")


# ══════════════════════════════════════════════════════════════════
# 6. 주주간계약서 (PRO)
# ══════════════════════════════════════════════════════════════════
def make_sha():
    doc = Document()
    page_setup(doc)

    title_para(doc, "주주간계약서", 18)
    title_para(doc, "Shareholders Agreement", 11)
    doc.add_paragraph()

    body(doc, "본 계약은 아래 주주들이 회사 경영 및 주식에 관한 사항을 정하기 위하여 체결한다.")
    doc.add_paragraph()

    heading(doc, "■ 회사 정보")
    field_line(doc, "회사명")
    field_line(doc, "사업자등록번호")
    field_line(doc, "주소")
    field_line(doc, "설립일")
    doc.add_paragraph()

    heading(doc, "■ 주주 현황")
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["주주명", "지분율 (%)", "주식 수", "주식 종류"]):
        table.rows[0].cells[i].text = h
        for r in table.rows[0].cells[i].paragraphs[0].runs:
            set_font(r, size=10, bold=True)
    doc.add_paragraph()

    heading(doc, "■ 제1조 (경영 참여)")
    body(doc, "이사회는 총 ___ 명으로 구성하며, 각 주주의 이사 지명권은 다음과 같다.")
    body(doc, "대주주 ( _______% 이상): ___ 명 지명 / 소수주주: ___ 명 지명")
    doc.add_paragraph()

    heading(doc, "■ 제2조 (주요 의사결정 사항)")
    body(doc, "다음 사항은 주주 전원의 동의 또는 지분 ___% 이상의 찬성을 요한다:")
    items = ["정관 변경", "자본금 증·감", "합병·분할·해산", "신주 발행 및 전환사채 발행",
             "자산의 처분 (총자산의 ___% 초과)", "대표이사 선임·해임"]
    for i, item in enumerate(items):
        body(doc, f"  {'①②③④⑤⑥'[i]} {item}")
    doc.add_paragraph()

    heading(doc, "■ 제3조 (우선 매수권)")
    body(doc, "주주가 주식을 제3자에게 양도하려는 경우, 다른 주주는 동일 조건으로 우선 매수할 권리를 갖는다.")
    body(doc, "우선 매수 의사 통보 기간: 양도 통지 수령 후  ___  일 이내")
    doc.add_paragraph()

    heading(doc, "■ 제4조 (공동 매도권 및 동반 매도 청구권)")
    body(doc, "대주주가 주식을 제3자에게 매도할 경우, 소수주주는 동일 조건으로 공동 매도에 참여할 수 있다.")
    body(doc, "대주주는 소수주주에게 동반 매도를 청구할 수 있다 (동반 매도 청구권, Drag-along).")
    doc.add_paragraph()

    heading(doc, "■ 제5조 (비희석 조항)")
    body(doc, "신주 발행 시 기존 주주는 지분 비율에 따라 신주를 우선 인수할 권리를 갖는다.")
    doc.add_paragraph()

    heading(doc, "■ 제6조 (배당)")
    body(doc, "배당은 이사회 결의에 의하며, 주주 간 지분 비율에 따라 지급한다.")
    doc.add_paragraph()

    heading(doc, "■ 제7조 (경업금지)")
    body(doc, "주주는 재임 중 및 퇴임 후 ___년간 회사와 동종 사업을 영위하거나 경쟁사에 참여할 수 없다.")
    doc.add_paragraph()

    heading(doc, "■ 제8조 (비밀유지)")
    body(doc, "주주는 회사의 영업 비밀 및 본 계약의 내용을 제3자에게 공개하여서는 아니 된다.")
    doc.add_paragraph()

    heading(doc, "■ 제9조 (계약 기간 및 해지)")
    body(doc, "본 계약은 체결일로부터 ___년간 유효하며, 당사자 전원의 합의로 변경 또는 해지할 수 있다.")
    doc.add_paragraph()

    heading(doc, "■ 제10조 (준거법 및 분쟁 해결)")
    body(doc, "본 계약은 대한민국 법률에 따르며, 분쟁은 대한상사중재원의 중재로 최종 해결한다.")
    doc.add_paragraph()

    divider(doc)
    body(doc, "계약 체결일:  ____년 __월 __일")
    doc.add_paragraph()
    sig_table(doc, ["주주 1 (갑)", "주주 2 (을)", "주주 3 (병)"])
    doc.add_paragraph()
    watermark_note(doc)

    doc.save(OUTPUT + "06_주주간계약서.docx")
    print("✓ 06_주주간계약서.docx")


# ══════════════════════════════════════════════════════════════════
# 7. 투자유치 텀시트
# ══════════════════════════════════════════════════════════════════
def make_termsheet():
    doc = Document()
    page_setup(doc)

    title_para(doc, "투자유치 텀시트", 18)
    title_para(doc, "Term Sheet (투자 조건 요약서)", 11)
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("※ 본 텀시트는 법적 구속력이 없는 투자 조건 요약서이며, 최종 계약서 체결 전 이해 당사자 간 협의를 위한 문서입니다.")
    set_font(run, size=9, color=(150, 0, 0))
    doc.add_paragraph()

    heading(doc, "■ 1. 기본 정보")
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("회사명", ""),
        ("투자 라운드", "□ Pre-Seed   □ Seed   □ Series A   □ Series B   □ 기타"),
        ("투자 금액", "금  _______________  원  (USD  _______________)"),
        ("투자 방식", "□ 보통주   □ 우선주   □ 전환사채(CB)   □ 신주인수권부사채(BW)"),
        ("Pre-money 밸류에이션", "금  _______________  원"),
        ("투자 후 지분율", "___  %"),
    ]
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for p2 in cell.paragraphs:
                for r in p2.runs:
                    set_font(r, size=10)
    doc.add_paragraph()

    heading(doc, "■ 2. 투자자 정보")
    field_line(doc, "리드 투자자")
    field_line(doc, "공동 투자자")
    field_line(doc, "담당자 / 연락처")
    doc.add_paragraph()

    heading(doc, "■ 3. 우선주 조건 (해당 시)")
    table2 = doc.add_table(rows=5, cols=2)
    table2.style = "Table Grid"
    rows2 = [
        ("우선 배당", "연  ___  %  누적  □  비누적  □"),
        ("청산 우선권", "___배  /  참가형  □  비참가형  □"),
        ("전환 비율", "1:1  (조정 조건: _______________)"),
        ("희석 방지 조항", "□ Full Ratchet   □ Weighted Average   □ 없음"),
        ("의결권", "□ 우선주 = 보통주   □ 별도 거부권 부여"),
    ]
    for i, (k, v) in enumerate(rows2):
        table2.rows[i].cells[0].text = k
        table2.rows[i].cells[1].text = v
        for cell in table2.rows[i].cells:
            for p2 in cell.paragraphs:
                for r in p2.runs:
                    set_font(r, size=10)
    doc.add_paragraph()

    heading(doc, "■ 4. 투자자 권리")
    items4 = [
        ("우선 매수권 (ROFR)", "□ 있음   □ 없음"),
        ("공동 매도권 (Tag-along)", "□ 있음   □ 없음"),
        ("동반 매도 청구권 (Drag-along)", "□ 있음   □ 없음"),
        ("이사회 참관권", "□ 있음   □ 없음  /  이사 지명권: ___명"),
        ("정보 열람권", "□ 분기  □ 반기  □ 연간 재무제표 제공"),
    ]
    table3 = doc.add_table(rows=len(items4), cols=2)
    table3.style = "Table Grid"
    for i, (k, v) in enumerate(items4):
        table3.rows[i].cells[0].text = k
        table3.rows[i].cells[1].text = v
        for cell in table3.rows[i].cells:
            for p2 in cell.paragraphs:
                for r in p2.runs:
                    set_font(r, size=10)
    doc.add_paragraph()

    heading(doc, "■ 5. 창업자 의무")
    body(doc, "경업 금지 기간:  투자 후  ___  년  /  지역: 대한민국 및 ( ___________ )")
    body(doc, "Lock-up 기간 (창업자 주식 매각 제한):  투자 후  ___  개월")
    body(doc, "전임 의무:  □ 있음   □ 없음  (위반 시 가중 희석 조항 적용)")
    doc.add_paragraph()

    heading(doc, "■ 6. EXIT 조건")
    body(doc, "IPO 목표 시기:  투자 후  ___  년 이내")
    body(doc, "M&A 시 투자자 동의 요부:  □ 필요   □ 불필요")
    body(doc, "풋옵션 (투자자 주식 매수 청구권):  □ 있음  ( _____년 후, ___배 수익 보장 )   □ 없음")
    doc.add_paragraph()

    heading(doc, "■ 7. 일정")
    body(doc, "텀시트 유효 기간:  ____년 __월 __일  까지")
    body(doc, "실사 (Due Diligence) 기간:  약  ___  주")
    body(doc, "최종 계약 목표일:  ____년 __월 __일")
    doc.add_paragraph()

    heading(doc, "■ 8. 기타 특약")
    for _ in range(3):
        body(doc, "• " + "_" * 65)
    doc.add_paragraph()

    divider(doc)
    body(doc, "작성일:  ____년 __월 __일")
    doc.add_paragraph()
    sig_table(doc, ["투자자", "피투자회사 대표"])
    doc.add_paragraph()
    watermark_note(doc)

    doc.save(OUTPUT + "07_투자유치_텀시트.docx")
    print("✓ 07_투자유치_텀시트.docx")


# ── 전체 실행 ──────────────────────────────────────────────────
if __name__ == "__main__":
    make_labor_contract()
    make_lease_contract()
    make_loan()
    make_nda()
    make_business_plan()
    make_sha()
    make_termsheet()
    print("\n✅ 전체 7개 서식 생성 완료 →", OUTPUT)
